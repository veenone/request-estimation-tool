"""Word report generation with python-docx.

Generates a formal estimation document per SPEC §5.2:
  1. Cover page with project name, estimation ID, request number, date, author
  2. Request details section
  3. Executive summary (total effort, feasibility, key risks)
  4. Project parameters table
  5. Detailed task breakdown table
  6. Tester allocation summary
  7. Timeline feasibility analysis
  8. Reference project comparison
  9. Assumptions and risk flags
"""

from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from .excel_report import ExcelReportData


def _add_styled_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers), style="Light Grid Accent 1")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)

    # Data rows
    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data):
            row.cells[i].text = str(value)
            for paragraph in row.cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)


def _feasibility_color(status: str) -> RGBColor:
    if status == "FEASIBLE":
        return RGBColor(0x00, 0x80, 0x00)
    elif status == "AT_RISK":
        return RGBColor(0xFF, 0xA5, 0x00)
    return RGBColor(0xFF, 0x00, 0x00)


def generate_word_report(data: ExcelReportData, output_path: str | Path | None = None) -> bytes | None:
    """Generate a Word report.

    Args:
        data: ExcelReportData with all estimation information.
        output_path: If provided, saves to file. If None, returns bytes.
    """
    doc = Document()

    # ── 1. Cover page ──────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Test Effort Estimation Report")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(data.project_name)
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    doc.add_paragraph()

    meta_items = [
        f"Estimation: {data.estimation_number}",
        f"Project Type: {data.project_type}",
    ]
    if data.request_number:
        meta_items.append(f"Request: {data.request_number}")
    meta_items.extend([
        f"Author: {data.created_by}",
        f"Date: {data.created_at}",
    ])

    for item in meta_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(item)
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    doc.add_page_break()

    # ── 2. Request details ─────────────────────────────
    if data.request_number:
        doc.add_heading("Request Details", level=1)
        _add_styled_table(doc, ["Field", "Value"], [
            ["Request Number", data.request_number],
            ["Title", data.project_name],
            ["Requester", data.requester_name],
            ["Business Unit", data.business_unit],
            ["Priority", data.priority],
            ["Requested Delivery", data.expected_delivery],
        ])
        doc.add_paragraph()

    # ── 3. Executive summary ───────────────────────────
    doc.add_heading("Executive Summary", level=1)

    summary_p = doc.add_paragraph()
    summary_p.add_run(f"This estimation covers the ").font.size = Pt(10)
    run = summary_p.add_run(data.project_type)
    run.bold = True
    run.font.size = Pt(10)
    summary_p.add_run(f" project \"{data.project_name}\" ").font.size = Pt(10)
    summary_p.add_run(f"with {data.dut_count} DUT(s), {data.profile_count} test profile(s), ").font.size = Pt(10)
    summary_p.add_run(f"and {data.pr_fix_count} PR fix(es) to validate.").font.size = Pt(10)

    doc.add_paragraph()

    _add_styled_table(doc, ["Metric", "Value"], [
        ["Grand Total (Hours)", f"{data.grand_total_hours:.1f}"],
        ["Grand Total (Person-Days)", f"{data.grand_total_days:.1f}"],
        ["Available Capacity (Hours)", f"{data.capacity_hours:.1f}"],
        ["Utilization", f"{data.utilization_pct:.1f}%"],
        ["Feasibility Status", data.feasibility_status],
    ])

    # Feasibility callout
    doc.add_paragraph()
    feasibility_p = doc.add_paragraph()
    run = feasibility_p.add_run(f"Feasibility: {data.feasibility_status}")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = _feasibility_color(data.feasibility_status)

    doc.add_paragraph()

    # ── 4. Project parameters ──────────────────────────
    doc.add_heading("Project Parameters", level=1)
    _add_styled_table(doc, ["Parameter", "Value"], [
        ["Project Type", data.project_type],
        ["DUT Count", str(data.dut_count)],
        ["Profile Count", str(data.profile_count)],
        ["DUT × Profile Combinations", str(data.dut_profile_combinations)],
        ["PR Fix Count", str(data.pr_fix_count)],
        ["Team Size", str(data.team_size)],
        ["Test Leader", "Yes" if data.has_leader else "No"],
        ["Expected Delivery", data.expected_delivery],
    ])

    doc.add_paragraph()

    # ── 4b. DUT x Profile Matrix ─────────────────────
    if data.dut_types and data.profiles:
        doc.add_heading("DUT x Profile Matrix", level=1)
        matrix_headers = ["DUT \\ Profile"] + [p.get("name", "") for p in data.profiles]
        active_combos = set()
        for combo in data.dut_profile_matrix:
            if len(combo) >= 2:
                active_combos.add((combo[0], combo[1]))
        matrix_rows = []
        for dut in data.dut_types:
            row = [dut.get("name", "")]
            for prof in data.profiles:
                is_active = (dut.get("id"), prof.get("id")) in active_combos or not data.dut_profile_matrix
                row.append("Yes" if is_active else "-")
            matrix_rows.append(row)
        _add_styled_table(doc, matrix_headers, matrix_rows)
        doc.add_paragraph()

    # ── 5. Task breakdown ──────────────────────────────
    doc.add_heading("Detailed Task Breakdown", level=1)

    task_headers = ["Task Name", "Type", "Base Hrs", "Tester Hrs", "Leader Hrs"]
    task_rows = []
    for task in data.tasks:
        task_rows.append([
            task.get("task_name", task.get("name", "")),
            task.get("task_type", ""),
            f"{task.get('base_hours', 0):.1f}",
            f"{task.get('calculated_hours', 0):.1f}",
            f"{task.get('leader_hours', 0):.1f}",
        ])
    _add_styled_table(doc, task_headers, task_rows)

    doc.add_paragraph()

    # ── 6. Effort summary ──────────────────────────────
    doc.add_heading("Effort Summary", level=1)
    _add_styled_table(doc, ["Component", "Hours"], [
        ["Total Tester Effort", f"{data.total_tester_hours:.1f}"],
        ["Test Leader Effort", f"{data.total_leader_hours:.1f}"],
        ["PR Fix Validation", f"{data.pr_fix_hours:.1f}"],
        ["New Feature Study", f"{data.study_hours:.1f}"],
        ["Buffer (10%)", f"{data.buffer_hours:.1f}"],
        ["GRAND TOTAL", f"{data.grand_total_hours:.1f}"],
    ])

    doc.add_paragraph()

    # ── 7. Feasibility analysis ────────────────────────
    doc.add_heading("Timeline Feasibility Analysis", level=1)

    p = doc.add_paragraph()
    p.add_run(f"With a team of {data.team_size} tester(s)").font.size = Pt(10)
    if data.has_leader:
        p.add_run(" and 1 test leader").font.size = Pt(10)
    p.add_run(f", the available capacity is {data.capacity_hours:.0f} hours.").font.size = Pt(10)

    p2 = doc.add_paragraph()
    p2.add_run(f"The estimated effort of {data.grand_total_hours:.1f} hours represents ").font.size = Pt(10)
    run = p2.add_run(f"{data.utilization_pct:.1f}%")
    run.bold = True
    run.font.size = Pt(10)
    p2.add_run(f" utilization, which is classified as ").font.size = Pt(10)
    run = p2.add_run(data.feasibility_status)
    run.bold = True
    run.font.color.rgb = _feasibility_color(data.feasibility_status)
    run.font.size = Pt(10)
    p2.add_run(".").font.size = Pt(10)

    doc.add_paragraph()

    # ── 8. Reference projects ──────────────────────────
    if data.reference_projects:
        doc.add_heading("Reference Project Comparison", level=1)
        ref_rows = []
        for proj in data.reference_projects:
            est = proj.get("estimated_hours", 0) or 0
            act = proj.get("actual_hours", 0) or 0
            ratio = f"{act/est:.2f}" if est > 0 else "N/A"
            ref_rows.append([
                proj.get("project_name", ""),
                proj.get("project_type", ""),
                f"{est:.0f}",
                f"{act:.0f}",
                ratio,
            ])
        _add_styled_table(doc, ["Project", "Type", "Estimated", "Actual", "Ratio"], ref_rows)
        doc.add_paragraph()

    # ── 9. Risk flags ──────────────────────────────────
    if data.risk_messages:
        doc.add_heading("Risk Flags and Assumptions", level=1)
        for msg in data.risk_messages:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(msg)
            run.font.size = Pt(10)

    # Save or return bytes
    if output_path:
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        return None
    else:
        buffer = BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
